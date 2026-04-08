import os
from docx import Document
from docx.parts.image import ImagePart
import zipfile
import re

def extract_formulas_and_images(doc_path, output_folder):
    doc = Document(doc_path)
    
    os.makedirs(output_folder, exist_ok=True)
    
    results = []
    formula_count = 0
    image_count = 0
    
    print(f"\n{'='*60}")
    print(f"处理文件: {doc_path}")
    print(f"{'='*60}")
    
    # 1. 提取嵌入的公式 (OMML - Office Math Markup Language)
    # Word中的公式通常存储在 document.xml 中的 o:math 或 m:oMath 标签
    print("\n[1] 检查嵌入公式...")
    
    # 读取document.xml查找公式
    doc_xml_path = 'word/document.xml'
    with zipfile.ZipFile(doc_path, 'r') as docx_file:
        try:
            doc_xml_content = docx_file.read(doc_xml_path).decode('utf-8')
            
            # 查找 Office Math 标签
            math_patterns = [
                r'<o:math>(.*?)</o:math>',
                r'<m:oMath>(.*?)</m:oMath>',
                r'<w:math>(.*?)</w:math>'
            ]
            
            for pattern in math_patterns:
                matches = re.findall(pattern, doc_xml_content, re.DOTALL)
                if matches:
                    print(f"  找到 {len(matches)} 个嵌入公式 (模式: {pattern[:20]}...)")
                    for i, match in enumerate(matches[:5]):  # 只显示前5个
                        formula_count += 1
                        print(f"    公式 {formula_count}: {match[:100]}...")
            
        except Exception as e:
            print(f"  读取document.xml失败: {e}")
    
    # 2. 提取图片（可能是公式图片）
    print("\n[2] 检查图片...")
    
    image_extensions = ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.emf', '.wmf']
    image_count = 0
    
    with zipfile.ZipFile(doc_path, 'r') as docx_file:
        for name in docx_file.namelist():
            # 检查是否是图片
            if any(name.lower().endswith(ext) for ext in image_extensions):
                image_count += 1
                # 提取图片
                image_data = docx_file.read(name)
                
                # 生成图片文件名
                img_filename = f"image_{image_count:03d}_{os.path.basename(name)}"
                img_path = os.path.join(output_folder, img_filename)
                
                with open(img_path, 'wb') as img_file:
                    img_file.write(image_data)
                
                print(f"    图片 {image_count}: {name} -> {img_path}")
    
    # 3. 遍历文档中的所有元素
    print("\n[3] 遍历文档段落和表格...")
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            # 检查是否包含公式相关关键词
            if any(keyword in text.lower() for keyword in ['公式', 'equation', 'model', '=', '∑', '∫', '∂', 'λ', 'β', 'α', 'γ']):
                results.append({
                    'type': 'paragraph',
                    'index': i,
                    'content': text[:200]
                })
                print(f"  段落 {i}: {text[:100]}...")
    
    # 4. 检查表格
    print("\n[4] 检查表格...")
    
    table_count = 0
    for table in doc.tables:
        table_count += 1
        print(f"  表格 {table_count}: {len(table.rows)} 行 x {len(table.columns)} 列")
    
    # 5. 提取文本内容用于分析
    print("\n[5] 完整文本内容预览...")
    
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # 保存文本内容
    text_output_path = os.path.join(output_folder, 'extracted_text.txt')
    with open(text_output_path, 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(full_text))
    print(f"  文本内容已保存到: {text_output_path}")
    
    print(f"\n{'='*60}")
    print(f"统计结果:")
    print(f"  - 嵌入公式: {formula_count}")
    print(f"  - 图片: {image_count}")
    print(f"  - 表格: {table_count}")
    print(f"  - 段落: {len(doc.paragraphs)}")
    print(f"{'='*60}")
    
    return {
        'formulas': formula_count,
        'images': image_count,
        'tables': table_count,
        'paragraphs': len(doc.paragraphs),
        'text_path': text_output_path
    }

if __name__ == "__main__":
    # 处理两个文档
    files = [
        r"E:\统计建模论文\论文\原版\1.0统计建模.docx",
        r"E:\统计建模论文\论文\原版\统计建模.docx"
    ]
    
    for doc_path in files:
        if os.path.exists(doc_path):
            output_folder = os.path.join(os.path.dirname(doc_path), 'extracted_' + os.path.basename(doc_path).replace('.docx', ''))
            results = extract_formulas_and_images(doc_path, output_folder)
        else:
            print(f"文件不存在: {doc_path}")
